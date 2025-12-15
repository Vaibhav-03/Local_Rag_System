"""
Retrieval System
================

Vector-based document retrieval using FAISS for efficient similarity search.

Why FAISS?
----------
- Highly optimized for similarity search
- Supports both exact and approximate nearest neighbor search
- Works well on CPU (no GPU required)
- Scales to millions of documents
- Memory-efficient index structures

Retrieval Flow:
1. Query → Embed query using same model as documents
2. Search FAISS index for nearest neighbors
3. Filter by similarity threshold
4. Return ranked documents with sources
"""

import os
import pickle
from typing import List, Optional, Tuple, Dict, Any
from pathlib import Path
import numpy as np

from .config import RetrieverConfig
from .embeddings import EmbeddingModel, DocumentChunker


class Document:
    """Represents a document chunk with metadata."""
    
    def __init__(
        self,
        text: str,
        doc_id: int,
        chunk_id: int = 0,
        source: str = "",
        metadata: Optional[Dict[str, Any]] = None,
    ):
        self.text = text
        self.doc_id = doc_id
        self.chunk_id = chunk_id
        self.source = source
        self.metadata = metadata or {}
    
    def __repr__(self):
        preview = self.text[:50] + "..." if len(self.text) > 50 else self.text
        return f"Document(doc_id={self.doc_id}, chunk={self.chunk_id}, text='{preview}')"
    
    def to_dict(self) -> dict:
        return {
            "text": self.text,
            "doc_id": self.doc_id,
            "chunk_id": self.chunk_id,
            "source": self.source,
            "metadata": self.metadata,
        }
    
    @classmethod
    def from_dict(cls, data: dict) -> "Document":
        return cls(**data)


class RetrievalResult:
    """Represents a retrieval result with score."""
    
    def __init__(self, document: Document, score: float, rank: int):
        self.document = document
        self.score = score
        self.rank = rank
    
    def __repr__(self):
        return f"RetrievalResult(rank={self.rank}, score={self.score:.4f}, doc={self.document})"
    
    def format_for_context(self) -> str:
        """Format the result for inclusion in LLM context."""
        source_info = f"[Source: {self.document.source}]" if self.document.source else f"[Chunk {self.document.chunk_id}]"
        return f"{source_info}\n{self.document.text}"


class VectorRetriever:
    """
    FAISS-based vector retriever for efficient document search.
    """
    
    def __init__(
        self,
        config: RetrieverConfig,
        embedding_model: EmbeddingModel,
    ):
        """
        Initialize the retriever.
        
        Args:
            config: Retriever configuration
            embedding_model: Embedding model for encoding queries/documents
        """
        self.config = config
        self.embedding_model = embedding_model
        self.index = None
        self.documents: List[Document] = []
        self.chunker = DocumentChunker(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
        )
        
        # Try to load existing index
        self._load_index()
    
    def _load_index(self) -> bool:
        """Load existing FAISS index and documents if available."""
        import faiss
        
        index_path = Path(self.config.index_path)
        docs_path = Path(self.config.documents_path)
        
        if index_path.exists() and docs_path.exists():
            try:
                print(f"Loading existing index from {index_path}")
                self.index = faiss.read_index(str(index_path))
                
                with open(docs_path, 'rb') as f:
                    self.documents = pickle.load(f)
                
                print(f"Loaded {len(self.documents)} documents")
                return True
            except Exception as e:
                print(f"Error loading index: {e}")
        
        return False
    
    def save_index(self) -> None:
        """Save the FAISS index and documents to disk."""
        import faiss
        
        if self.index is None:
            raise RuntimeError("No index to save")
        
        index_path = Path(self.config.index_path)
        docs_path = Path(self.config.documents_path)
        
        # Create directories if needed
        index_path.parent.mkdir(parents=True, exist_ok=True)
        docs_path.parent.mkdir(parents=True, exist_ok=True)
        
        print(f"Saving index to {index_path}")
        faiss.write_index(self.index, str(index_path))
        
        with open(docs_path, 'wb') as f:
            pickle.dump(self.documents, f)
        
        print(f"Saved {len(self.documents)} documents")
    
    def build_index(self, documents: List[Document]) -> None:
        """
        Build a new FAISS index from documents.
        
        Args:
            documents: List of Document objects to index
        """
        import faiss
        
        if not documents:
            raise ValueError("No documents provided")
        
        print(f"Building index for {len(documents)} documents...")
        self.documents = documents
        
        # Encode all documents
        texts = [doc.text for doc in documents]
        embeddings = self.embedding_model.encode_documents(texts, show_progress=True)
        
        # Create FAISS index
        dimension = embeddings.shape[1]
        
        # Use IndexFlatIP for inner product (cosine similarity with normalized vectors)
        self.index = faiss.IndexFlatIP(dimension)
        
        # Add embeddings to index
        self.index.add(embeddings.astype(np.float32))
        
        print(f"Index built with {self.index.ntotal} vectors")
    
    def add_documents(self, documents: List[Document]) -> None:
        """
        Add new documents to existing index.
        
        Args:
            documents: New documents to add
        """
        import faiss
        
        if not documents:
            return
        
        # Encode new documents
        texts = [doc.text for doc in documents]
        embeddings = self.embedding_model.encode_documents(texts, show_progress=True)
        
        if self.index is None:
            # Create new index
            dimension = embeddings.shape[1]
            self.index = faiss.IndexFlatIP(dimension)
        
        # Add to index
        self.index.add(embeddings.astype(np.float32))
        self.documents.extend(documents)
        
        print(f"Added {len(documents)} documents. Total: {len(self.documents)}")
    
    def retrieve(
        self,
        query: str,
        top_k: Optional[int] = None,
        threshold: Optional[float] = None,
    ) -> List[RetrievalResult]:
        """
        Retrieve relevant documents for a query.
        
        Args:
            query: Search query
            top_k: Number of results to return (default from config)
            threshold: Minimum similarity threshold (default from config)
            
        Returns:
            List of RetrievalResult objects, ranked by relevance
        """
        if self.index is None or self.index.ntotal == 0:
            return []
        
        top_k = top_k or self.config.top_k
        threshold = threshold or self.config.similarity_threshold
        
        # Encode query
        query_embedding = self.embedding_model.encode_query(query)
        
        # Search index
        scores, indices = self.index.search(
            query_embedding.reshape(1, -1).astype(np.float32),
            min(top_k, self.index.ntotal),
        )
        
        # Build results
        results = []
        for rank, (score, idx) in enumerate(zip(scores[0], indices[0])):
            if idx == -1:  # FAISS returns -1 for missing results
                continue
            if score < threshold:
                continue
            
            document = self.documents[idx]
            results.append(RetrievalResult(
                document=document,
                score=float(score),
                rank=rank + 1,
            ))
        
        return results
    
    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the index."""
        return {
            "num_documents": len(self.documents),
            "index_size": self.index.ntotal if self.index else 0,
            "dimension": self.embedding_model.dimension,
        }


def load_corpus_from_directory(
    directory: str,
    chunker: DocumentChunker,
) -> List[Document]:
    """
    Load documents from a directory of text files.
    
    Supports: .txt, .md, .pdf, .docx
    
    Args:
        directory: Path to corpus directory
        chunker: Document chunker for splitting
        
    Returns:
        List of Document objects
    """
    from pathlib import Path
    
    corpus_path = Path(directory)
    if not corpus_path.exists():
        print(f"Corpus directory not found: {directory}")
        return []
    
    documents = []
    doc_id = 0
    
    # Supported file types
    text_extensions = {'.txt', '.md'}
    
    for file_path in corpus_path.rglob('*'):
        if not file_path.is_file():
            continue
        
        ext = file_path.suffix.lower()
        
        try:
            if ext in text_extensions:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif ext == '.pdf':
                text = _load_pdf(file_path)
            elif ext == '.docx':
                text = _load_docx(file_path)
            else:
                continue
            
            if not text.strip():
                continue
            
            # Chunk the document
            source_name = str(file_path.relative_to(corpus_path))
            chunks = chunker.chunk_text(text, {"source_file": source_name})
            
            for chunk in chunks:
                documents.append(Document(
                    text=chunk["text"],
                    doc_id=doc_id,
                    chunk_id=chunk["chunk_id"],
                    source=source_name,
                    metadata=chunk.get("metadata", {}),
                ))
            
            doc_id += 1
            print(f"Loaded: {source_name} ({len(chunks)} chunks)")
            
        except Exception as e:
            print(f"Error loading {file_path}: {e}")
    
    print(f"Total: {len(documents)} document chunks from {doc_id} files")
    return documents


def _load_pdf(file_path: Path) -> str:
    """Load text from PDF file."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except ImportError:
        print("pypdf not installed. Skipping PDF files.")
        return ""


def _load_docx(file_path: Path) -> str:
    """Load text from DOCX file."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except ImportError:
        print("python-docx not installed. Skipping DOCX files.")
        return ""

